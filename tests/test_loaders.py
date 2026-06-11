"""Loader和SmartSplitter模块测试"""
import tempfile
from pathlib import Path

import pytest

from src.core.loaders.base import DocumentElement, ElementType
from src.core.loaders.markdown_loader import MarkdownLoader
from src.core.loaders.docx_loader import DocxLoader
from src.core.loaders.txt_loader import TxtLoader
from src.core.splitter import SmartSplitter, Chunk


# ---------------------------------------------------------------------------
# MarkdownLoader 测试
# ---------------------------------------------------------------------------

class TestMarkdownLoader:
    """Markdown加载器测试"""

    def test_load_md_file(self, tmp_path: Path):
        """能加载 .md 文件并返回正确内容"""
        md_file = tmp_path / "test.md"
        md_file.write_text("# Hello\n\n这是正文内容。", encoding="utf-8")

        loader = MarkdownLoader()
        elements = loader.load(md_file)

        assert len(elements) == 1
        assert elements[0].content == "# Hello\n\n这是正文内容。"
        assert elements[0].element_type == ElementType.TEXT
        assert elements[0].source_file == str(md_file)

    def test_load_empty_md_file(self, tmp_path: Path):
        """空文件返回空列表"""
        md_file = tmp_path / "empty.md"
        md_file.write_text("", encoding="utf-8")

        loader = MarkdownLoader()
        elements = loader.load(md_file)
        assert elements == []

    def test_load_nonexistent_file(self):
        """加载不存在的文件抛出 FileNotFoundError"""
        loader = MarkdownLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(Path("/nonexistent/file.md"))

    def test_can_handle(self):
        """can_handle 判断正确"""
        assert MarkdownLoader.can_handle(Path("a.md")) is True
        assert MarkdownLoader.can_handle(Path("a.markdown")) is True
        assert MarkdownLoader.can_handle(Path("a.txt")) is False

    def test_supported_extensions(self):
        """支持的扩展名正确"""
        assert ".md" in MarkdownLoader.supported_extensions()
        assert ".markdown" in MarkdownLoader.supported_extensions()


# ---------------------------------------------------------------------------
# DocxLoader 测试
# ---------------------------------------------------------------------------

class TestDocxLoader:
    """Word文档加载器测试"""

    def _create_simple_docx(self, path: Path, paragraphs: list[str], table_data: list[list[str]] | None = None):
        """创建一个简单的测试docx文件"""
        from docx import Document

        doc = Document()
        for text in paragraphs:
            doc.add_paragraph(text)

        if table_data:
            table = doc.add_table(rows=len(table_data), cols=len(table_data[0]))
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data):
                    table.rows[i].cells[j].text = cell_text

        doc.save(str(path))

    def test_load_docx_paragraphs(self, tmp_path: Path):
        """能加载docx中的段落"""
        docx_file = tmp_path / "test.docx"
        self._create_simple_docx(docx_file, ["第一段", "第二段"])

        loader = DocxLoader()
        elements = loader.load(docx_file)

        assert len(elements) == 2
        assert elements[0].content == "第一段"
        assert elements[0].element_type == ElementType.TEXT
        assert elements[1].content == "第二段"

    def test_load_docx_with_table(self, tmp_path: Path):
        """能加载docx中的表格"""
        docx_file = tmp_path / "table.docx"
        table_data = [["姓名", "年龄"], ["张三", "30"], ["李四", "25"]]
        self._create_simple_docx(docx_file, ["以下是表格："], table_data)

        loader = DocxLoader()
        elements = loader.load(docx_file)

        # 至少有段落和表格
        assert len(elements) >= 2
        table_elements = [e for e in elements if e.element_type == ElementType.TABLE]
        assert len(table_elements) == 1
        # 表格内容应为 Markdown 格式
        table_content = table_elements[0].content
        assert "姓名" in table_content
        assert "张三" in table_content
        assert "---" in table_content  # Markdown 分隔行

    def test_load_empty_docx(self, tmp_path: Path):
        """空docx文件返回空列表"""
        from docx import Document

        docx_file = tmp_path / "empty.docx"
        Document().save(str(docx_file))

        loader = DocxLoader()
        elements = loader.load(docx_file)
        assert elements == []

    def test_load_nonexistent_file(self):
        """加载不存在的文件抛出 FileNotFoundError"""
        loader = DocxLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(Path("/nonexistent/file.docx"))

    def test_can_handle(self):
        """can_handle 判断正确"""
        assert DocxLoader.can_handle(Path("a.docx")) is True
        assert DocxLoader.can_handle(Path("a.doc")) is True
        assert DocxLoader.can_handle(Path("a.md")) is False

    def test_supported_extensions(self):
        """支持的扩展名正确"""
        assert ".docx" in DocxLoader.supported_extensions()
        assert ".doc" in DocxLoader.supported_extensions()


# ---------------------------------------------------------------------------
# TxtLoader 测试
# ---------------------------------------------------------------------------

class TestTxtLoader:
    """纯文本加载器测试"""

    def test_load_txt_file(self, tmp_path: Path):
        """能加载 .txt 文件"""
        txt_file = tmp_path / "test.txt"
        txt_file.write_text("这是一段测试文本。\n第二行内容。", encoding="utf-8")

        loader = TxtLoader()
        elements = loader.load(txt_file)

        assert len(elements) == 1
        assert "测试文本" in elements[0].content
        assert elements[0].element_type == ElementType.TEXT
        assert elements[0].source_file == str(txt_file)

    def test_load_empty_txt_file(self, tmp_path: Path):
        """空文件返回空列表"""
        txt_file = tmp_path / "empty.txt"
        txt_file.write_text("", encoding="utf-8")

        loader = TxtLoader()
        elements = loader.load(txt_file)
        assert elements == []

    def test_load_nonexistent_file(self):
        """加载不存在的文件抛出 FileNotFoundError"""
        loader = TxtLoader()
        with pytest.raises(FileNotFoundError):
            loader.load(Path("/nonexistent/file.txt"))

    def test_can_handle(self):
        """can_handle 判断正确"""
        assert TxtLoader.can_handle(Path("a.txt")) is True
        assert TxtLoader.can_handle(Path("a.text")) is True
        assert TxtLoader.can_handle(Path("a.log")) is True
        assert TxtLoader.can_handle(Path("a.md")) is False

    def test_supported_extensions(self):
        """支持的扩展名正确"""
        assert ".txt" in TxtLoader.supported_extensions()
        assert ".text" in TxtLoader.supported_extensions()
        assert ".log" in TxtLoader.supported_extensions()


# ---------------------------------------------------------------------------
# SmartSplitter 测试
# ---------------------------------------------------------------------------

class TestSmartSplitter:
    """SmartSplitter 对不同 element_type 的处理测试"""

    def test_table_not_split(self):
        """表格类型不切分，整体保留"""
        splitter = SmartSplitter(chunk_size=100, chunk_overlap=10)
        table_content = "| 姓名 | 年龄 |\n| --- | --- |\n| 张三 | 30 |"
        element = DocumentElement(
            content=table_content,
            element_type=ElementType.TABLE,
            metadata={"rows": 2, "cols": 2},
            source_file="test.docx",
        )

        chunks = splitter.split_element(element)

        assert len(chunks) == 1
        assert chunks[0].content == table_content
        assert chunks[0].metadata["element_type"] == "table"

    def test_image_description_not_split(self):
        """图片描述类型不切分，整体保留"""
        splitter = SmartSplitter(chunk_size=100, chunk_overlap=10)
        desc = "这是一张包含统计图表的图片，展示了2024年的销售趋势。"
        element = DocumentElement(
            content=desc,
            element_type=ElementType.IMAGE_DESCRIPTION,
            metadata={},
            source_file="test.pdf",
        )

        chunks = splitter.split_element(element)

        assert len(chunks) == 1
        assert chunks[0].content == desc
        assert chunks[0].metadata["element_type"] == "image_description"

    def test_code_not_split(self):
        """代码块类型不切分，整体保留"""
        splitter = SmartSplitter(chunk_size=100, chunk_overlap=10)
        code = "def hello():\n    print('hello world')"
        element = DocumentElement(
            content=code,
            element_type=ElementType.CODE,
            metadata={},
            source_file="test.py",
        )

        chunks = splitter.split_element(element)

        assert len(chunks) == 1
        assert chunks[0].content == code
        assert chunks[0].metadata["element_type"] == "code"

    def test_text_markdown_style_split(self):
        """文本类型（Markdown风格）按标题切分"""
        splitter = SmartSplitter(chunk_size=500, chunk_overlap=50)
        content = "# 第一章\n这是第一章的内容。\n\n## 第二节\n这是第二节的内容。\n\n## 第三节\n这是第三节的内容。"
        element = DocumentElement(
            content=content,
            element_type=ElementType.TEXT,
            source_file="test.md",
        )

        chunks = splitter.split_element(element)

        assert len(chunks) == 3
        # 验证 section 名
        sections = [c.metadata.get("section", "") for c in chunks]
        assert "第二章" not in sections  # "第一行没有 ## 但整个 section header是"第一行的标题"
        # 验证内容被正确分割
        assert "第一章" in chunks[0].content
        assert "第二节" in chunks[1].content
        assert "第三节" in chunks[2].content

    def test_text_plain_style_split(self):
        """文本类型（纯文本风格）使用 RecursiveCharacterSplitter"""
        splitter = SmartSplitter(chunk_size=30, chunk_overlap=5)
        # 短文本，不含Markdown标记 -> MarkdownSplitter返回1个chunk
        # 但长度 < chunk_size // 2 = 15，所以走 text_splitter
        content = "这是纯文本内容"
        element = DocumentElement(
            content=content,
            element_type=ElementType.TEXT,
            source_file="test.txt",
        )

        chunks = splitter.split_element(element)
        # 短文本应该只有1个chunk
        assert len(chunks) >= 1
        assert chunks[0].content == content

    def test_split_elements_batch(self):
        """split_elements 批量处理多个元素"""
        splitter = SmartSplitter(chunk_size=200, chunk_overlap=20)

        elements = [
            DocumentElement(
                content="# 标题\n内容A",
                element_type=ElementType.TEXT,
                source_file="a.md",
            ),
            DocumentElement(
                content="| 列1 | 列2 |\n| --- | --- |\n| 值1 | 值2 |",
                element_type=ElementType.TABLE,
                source_file="b.docx",
            ),
        ]

        chunks = splitter.split_elements(elements)

        # 文本1个 + 表格1个 = 2个
        assert len(chunks) >= 2
        # 表格的 chunk 元数据 element_type 应为 table
        table_chunks = [c for c in chunks if c.metadata.get("element_type") == "table"]
        assert len(table_chunks) == 1

    def test_metadata_preserved_after_split(self):
        """切分后 metadata 正确保留"""
        splitter = SmartSplitter(chunk_size=500, chunk_overlap=50)
        element = DocumentElement(
            content="# 标题\n内容",
            element_type=ElementType.TEXT,
            metadata={"custom_key": "custom_value"},
            source_file="test.md",
            page_number=3,
        )

        chunks = splitter.split_element(element)

        assert len(chunks) >= 1
        assert chunks[0].metadata["source_file"] == "test.md"
        assert chunks[0].metadata["element_type"] == "text"
        assert chunks[0].metadata["page_number"] == 3
        assert chunks[0].metadata["custom_key"] == "custom_value"

    def test_code_block_inside_markdown_not_split(self):
        """Markdown中的代码块内不切分"""
        splitter = SmartSplitter(chunk_size=100, chunk_overlap=10)
        content = "# 标题\n\n这是正文\n\n```\ndef hello():\n    pass\n```"
        element = DocumentElement(
            content=content,
            element_type=ElementType.TEXT,
            source_file="test.md",
        )

        chunks = splitter.split_element(element)

        # 整体应该只有1个section（标题行开始），且包含完整代码块
        all_content = "\n".join(c.content for c in chunks)
        assert "```" in all_content
        assert "def hello():" in all_content
        assert "pass" in all_content
