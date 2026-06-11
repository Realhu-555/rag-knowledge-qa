"""Word文档加载器"""
from pathlib import Path

from src.core.loaders.base import BaseLoader, DocumentElement, ElementType


def _table_to_markdown(table) -> str:
    """将docx表格转换为Markdown格式

    Args:
        table: python-docx表格对象

    Returns:
        Markdown格式的表格字符串
    """
    rows = []
    for row in table.rows:
        cells = []
        for cell in row.cells:
            # 清理单元格文本，替换换行符
            cell_text = cell.text.replace('\n', ' ').strip()
            cells.append(cell_text)
        rows.append(cells)

    if not rows:
        return ""

    # 构建Markdown表格
    lines = []
    # 表头
    lines.append("| " + " | ".join(rows[0]) + " |")
    # 分隔符
    lines.append("| " + " | ".join(["---"] * len(rows[0])) + " |")
    # 数据行
    for row in rows[1:]:
        # 确保每行的列数一致
        while len(row) < len(rows[0]):
            row.append("")
        lines.append("| " + " | ".join(row[:len(rows[0])]) + " |")

    return "\n".join(lines)


class DocxLoader(BaseLoader):
    """Word文档加载器，提取段落和表格"""

    @classmethod
    def can_handle(cls, file_path: Path) -> bool:
        """判断是否能处理该文件

        Args:
            file_path: 文件路径

        Returns:
            True表示可以处理
        """
        return file_path.suffix.lower() in ['.docx', '.doc']

    def load(self, file_path: Path) -> list[DocumentElement]:
        """加载Word文档

        Args:
            file_path: Word文档路径

        Returns:
            文档元素列表
        """
        if not file_path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装python-docx: uv pip install python-docx")

        doc = Document(file_path)
        elements = []

        # 预建索引，避免每次查找都线性扫描
        para_index = {para._element: para for para in doc.paragraphs}
        table_index = {tbl._tbl: tbl for tbl in doc.tables}

        # 按文档顺序遍历段落和表格
        # python-docx的doc.element.body包含所有块级元素
        for child in doc.element.body:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag

            if tag == 'p':
                # 段落
                paragraph = para_index.get(child)
                if paragraph:
                    content = paragraph.text.strip()
                    if content:
                        elements.append(DocumentElement(
                            content=content,
                            element_type=ElementType.TEXT,
                            metadata={"element": "paragraph"},
                            source_file=str(file_path),
                        ))
            elif tag == 'tbl':
                # 表格
                table = table_index.get(child)
                if table:
                    markdown_table = _table_to_markdown(table)
                    if markdown_table:
                        elements.append(DocumentElement(
                            content=markdown_table,
                            element_type=ElementType.TABLE,
                            metadata={
                                "element": "table",
                                "rows": len(table.rows),
                                "cols": len(table.columns) if table.rows else 0,
                            },
                            source_file=str(file_path),
                        ))

        return elements

    @classmethod
    def supported_extensions(cls) -> list[str]:
        """返回支持的文件扩展名"""
        return ['.docx', '.doc']
